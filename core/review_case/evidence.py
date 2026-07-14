from __future__ import annotations

import fcntl
import hashlib
import io
import json
import os
import re
import subprocess
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable

from .identity import ROOT, find_work, validate_work_identity


DEFAULT_WORKS = ROOT / "config" / "works.yaml"


EVIDENCE_ID_RE = re.compile(r"^E\d{3,}$")
EVIDENCE_REF_RE = re.compile(r"\[@(E\d{3,})\]")
ALLOWED_EVIDENCE_KINDS = {
    "source_span",
    "document_span",
    "git_commit",
    "fingerprint_comparison",
    "search_result",
}


class EvidenceError(ValueError):
    pass


@dataclass(frozen=True)
class EvidenceCard:
    evidence_id: str
    kind: str
    title: str
    source: dict[str, Any]
    excerpt: str
    facts: list[dict[str, Any]]
    table: dict[str, Any] | None = None
    raw_ref: dict[str, str] | None = None

    def as_dict(self) -> dict[str, Any]:
        result: dict[str, Any] = {
            "evidence_id": self.evidence_id,
            "kind": self.kind,
            "title": self.title,
            "source": self.source,
            "excerpt": self.excerpt,
            "facts": self.facts,
        }
        if self.table:
            result["table"] = self.table
        if self.raw_ref:
            result["raw_ref"] = self.raw_ref
        return result


@dataclass(frozen=True)
class GitSource:
    work_id: str
    display_name: str
    repo: Path


def extract_refs(text: str) -> set[str]:
    return set(EVIDENCE_REF_RE.findall(text))


def _validate_card(data: dict[str, Any], *, line_no: int) -> EvidenceCard:
    evidence_id = str(data.get("evidence_id", "")).strip()
    if not EVIDENCE_ID_RE.fullmatch(evidence_id):
        raise EvidenceError(f"evidence.jsonl 第 {line_no} 行的 evidence_id 非法: {evidence_id}")

    kind = str(data.get("kind", "")).strip()
    if kind not in ALLOWED_EVIDENCE_KINDS:
        raise EvidenceError(f"{evidence_id} 的 kind 非法: {kind}")

    title = str(data.get("title", "")).strip()
    excerpt = str(data.get("excerpt", ""))
    source = data.get("source")
    facts = data.get("facts", [])
    table = data.get("table")
    raw_ref = data.get("raw_ref")

    if not title:
        raise EvidenceError(f"{evidence_id} 缺少 title")
    if not isinstance(source, dict) or not str(source.get("display_name", "")).strip():
        raise EvidenceError(f"{evidence_id} 的 source 缺少 display_name")
    if not isinstance(facts, list) or any(not isinstance(item, dict) for item in facts):
        raise EvidenceError(f"{evidence_id} 的 facts 必须是对象列表")
    if table is not None:
        if not isinstance(table, dict):
            raise EvidenceError(f"{evidence_id} 的 table 必须是对象")
        columns = table.get("columns")
        rows = table.get("rows")
        if not isinstance(columns, list) or not columns:
            raise EvidenceError(f"{evidence_id} 的 table.columns 不能为空")
        if not isinstance(rows, list) or any(not isinstance(row, list) for row in rows):
            raise EvidenceError(f"{evidence_id} 的 table.rows 必须是二维列表")
        if any(len(row) != len(columns) for row in rows):
            raise EvidenceError(f"{evidence_id} 的 table 行列数量不一致")
    if raw_ref is not None:
        if not isinstance(raw_ref, dict) or not raw_ref.get("path") or not raw_ref.get("sha256"):
            raise EvidenceError(f"{evidence_id} 的 raw_ref 必须包含 path 和 sha256")

    return EvidenceCard(
        evidence_id=evidence_id,
        kind=kind,
        title=title,
        source=source,
        excerpt=excerpt,
        facts=facts,
        table=table,
        raw_ref=raw_ref,
    )


def load_evidence(path: Path) -> dict[str, EvidenceCard]:
    cards: dict[str, EvidenceCard] = {}
    if not path.exists():
        return cards
    for line_no, line in enumerate(path.read_text(encoding="utf-8").splitlines(), start=1):
        if not line.strip():
            continue
        try:
            data = json.loads(line)
        except json.JSONDecodeError as exc:
            raise EvidenceError(f"evidence.jsonl 第 {line_no} 行不是合法 JSON: {exc}") from exc
        if not isinstance(data, dict):
            raise EvidenceError(f"evidence.jsonl 第 {line_no} 行必须是对象")
        card = _validate_card(data, line_no=line_no)
        if card.evidence_id in cards:
            raise EvidenceError(f"重复 evidence_id: {card.evidence_id}")
        cards[card.evidence_id] = card
    return cards


def _run_git(repo: Path, *args: str, text: bool = True) -> str | bytes:
    result = subprocess.run(
        ["git", "-C", str(repo), *args],
        check=False,
        capture_output=True,
        text=text,
    )
    if result.returncode != 0:
        stderr = result.stderr.strip() if text else result.stderr.decode("utf-8", errors="replace").strip()
        raise EvidenceError(stderr or f"git {' '.join(args)} 执行失败")
    return result.stdout


def _lock_commit(repo: Path, commit: str) -> str:
    value = str(_run_git(repo, "rev-parse", "--verify", f"{commit}^{{commit}}"))
    return value.strip()


def _resolve_source(
    *,
    work_id: str | None,
    repo: str | Path | None,
    display_name: str | None,
    external_id: str | None = None,
    works_path: Path = DEFAULT_WORKS,
) -> GitSource:
    if bool(work_id) == bool(repo):
        raise EvidenceError("必须且只能指定 --work-id 或 --repo")
    if work_id:
        work = find_work(work_id, works_path)
        if work is None:
            raise EvidenceError(f"works.yaml 中不存在 work_id: {work_id}")
        report = validate_work_identity(work)
        report.raise_for_errors()
        return GitSource(work.work_id, work.display_name, work.repo_path)

    repo_path = Path(str(repo)).expanduser().resolve()
    if not display_name:
        raise EvidenceError("使用 --repo 时必须指定 --display-name")
    if not repo_path.is_dir():
        raise EvidenceError(f"Git 仓库不存在: {repo_path}")
    _run_git(repo_path, "rev-parse", "--git-dir")
    return GitSource(external_id or repo_path.name, display_name, repo_path)


def _parse_line_range(value: str) -> tuple[int, int]:
    match = re.fullmatch(r"(\d+)(?::(\d+))?", value.strip())
    if not match:
        raise EvidenceError("行号必须使用 START 或 START:END")
    start = int(match.group(1))
    end = int(match.group(2) or start)
    if start < 1 or end < start:
        raise EvidenceError("行号范围非法")
    return start, end


def _select_lines(text: str, value: str) -> tuple[str, int, int]:
    start, end = _parse_line_range(value)
    lines = text.splitlines()
    if end > len(lines):
        raise EvidenceError(f"行号超出范围: 文件共 {len(lines)} 行")
    return "\n".join(lines[start - 1 : end]), start, end


def _sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _display_path(path: Path) -> str:
    try:
        return path.resolve().relative_to(ROOT).as_posix()
    except ValueError:
        return str(path.resolve())


def _dedup_signature(card: dict[str, Any]) -> str:
    source = card["source"]
    raw_ref = card.get("raw_ref") or {}
    payload: dict[str, Any] = {
        "kind": card["kind"],
        "source": source,
        "raw_sha256": raw_ref.get("sha256", ""),
    }
    if card["kind"] == "search_result" and not raw_ref:
        payload["facts"] = card.get("facts", [])
    encoded = json.dumps(payload, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    return hashlib.sha256(encoded.encode("utf-8")).hexdigest()


def add_evidence(
    case_dir: Path,
    *,
    kind: str,
    title: str,
    source: dict[str, Any],
    excerpt: str,
    facts: list[dict[str, Any]] | None = None,
    table: dict[str, Any] | None = None,
    raw_ref: dict[str, str] | None = None,
) -> str:
    if kind not in ALLOWED_EVIDENCE_KINDS:
        raise EvidenceError(f"不支持的 evidence kind: {kind}")
    case_dir.mkdir(parents=True, exist_ok=True)
    evidence_path = case_dir / "evidence.jsonl"
    lock_path = case_dir / "case_state" / "evidence.lock"
    lock_path.parent.mkdir(parents=True, exist_ok=True)

    draft: dict[str, Any] = {
        "kind": kind,
        "title": title.strip(),
        "source": source,
        "excerpt": excerpt,
        "facts": facts or [],
    }
    if table:
        draft["table"] = table
    if raw_ref:
        draft["raw_ref"] = raw_ref
    signature = _dedup_signature(draft)

    with lock_path.open("a+", encoding="utf-8") as lock_handle:
        fcntl.flock(lock_handle.fileno(), fcntl.LOCK_EX)
        cards = load_evidence(evidence_path)
        for existing in cards.values():
            if _dedup_signature(existing.as_dict()) == signature:
                return existing.evidence_id

        next_number = max((int(item[1:]) for item in cards), default=0) + 1
        evidence_id = f"E{next_number:03d}"
        payload = {"evidence_id": evidence_id, **draft}
        _validate_card(payload, line_no=next_number)
        with evidence_path.open("a", encoding="utf-8") as output:
            output.write(json.dumps(payload, ensure_ascii=False, sort_keys=True) + "\n")
            output.flush()
            os.fsync(output.fileno())
        return evidence_id


def capture_span(
    case_dir: Path,
    *,
    title: str,
    commit: str,
    path: str,
    lines: str,
    work_id: str | None = None,
    repo: str | Path | None = None,
    display_name: str | None = None,
    external_id: str | None = None,
    works_path: Path = DEFAULT_WORKS,
) -> str:
    source_spec = _resolve_source(
        work_id=work_id,
        repo=repo,
        display_name=display_name,
        external_id=external_id,
        works_path=works_path,
    )
    locked_commit = _lock_commit(source_spec.repo, commit)
    blob_hash = str(_run_git(source_spec.repo, "rev-parse", f"{locked_commit}:{path}")).strip()
    content = bytes(_run_git(source_spec.repo, "show", f"{locked_commit}:{path}", text=False)).decode(
        "utf-8", errors="replace"
    )
    excerpt, start, end = _select_lines(content, lines)
    return add_evidence(
        case_dir,
        kind="source_span",
        title=title,
        source={
            "work_id": source_spec.work_id,
            "display_name": source_spec.display_name,
            "commit": locked_commit,
            "path": path,
            "locator": f"L{start}-L{end}",
            "line_start": start,
            "line_end": end,
            "object_hash": blob_hash,
        },
        excerpt=excerpt,
        facts=[],
    )


def capture_document(
    case_dir: Path,
    *,
    title: str,
    commit: str | None = None,
    path: str | None = None,
    file_path: Path | None = None,
    page: int | None = None,
    paragraph: int | None = None,
    lines: str | None = None,
    work_id: str | None = None,
    repo: str | Path | None = None,
    display_name: str | None = None,
    external_id: str | None = None,
    works_path: Path = DEFAULT_WORKS,
) -> str:
    source_spec = _resolve_source(
        work_id=work_id,
        repo=repo,
        display_name=display_name,
        external_id=external_id,
        works_path=works_path,
    )
    raw_ref: dict[str, str] | None = None
    if file_path is not None:
        if commit or path:
            raise EvidenceError("--file 不能与 --commit/--path 同时使用")
        local_path = file_path.expanduser().resolve()
        if not local_path.is_file():
            raise EvidenceError(f"文档不存在: {local_path}")
        data = local_path.read_bytes()
        locked_commit = ""
        blob_hash = hashlib.sha256(data).hexdigest()
        document_path = local_path.name
        raw_ref = {"path": _display_path(local_path), "sha256": blob_hash}
    else:
        if not commit or not path:
            raise EvidenceError("Git 文档必须同时指定 --commit 和 --path")
        locked_commit = _lock_commit(source_spec.repo, commit)
        blob_hash = str(_run_git(source_spec.repo, "rev-parse", f"{locked_commit}:{path}")).strip()
        data = bytes(_run_git(source_spec.repo, "show", f"{locked_commit}:{path}", text=False))
        document_path = path
    suffix = Path(document_path).suffix.lower()
    source_extra: dict[str, Any] = {}

    if suffix == ".pdf":
        if page is None or page < 1:
            raise EvidenceError("PDF 证据必须指定从 1 开始的 --page")
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        if page > len(reader.pages):
            raise EvidenceError(f"PDF 页码超出范围: 共 {len(reader.pages)} 页")
        page_text = reader.pages[page - 1].extract_text() or ""
        if lines:
            excerpt, start, end = _select_lines(page_text, lines)
            locator = f"P{page}:L{start}-L{end}"
            source_extra.update({"line_start": start, "line_end": end})
        else:
            if len(page_text) > 12000:
                raise EvidenceError("PDF 页面文本过长，请使用 --lines 固定更小范围")
            excerpt = page_text
            locator = f"P{page}"
        source_extra.update({"page": page, "format": "pdf"})
    elif suffix == ".docx":
        if paragraph is None or paragraph < 1:
            raise EvidenceError("DOCX 证据必须指定从 1 开始的 --paragraph")
        from docx import Document

        document = Document(io.BytesIO(data))
        if paragraph > len(document.paragraphs):
            raise EvidenceError(f"DOCX 段落超出范围: 共 {len(document.paragraphs)} 段")
        excerpt = document.paragraphs[paragraph - 1].text
        locator = f"paragraph {paragraph}"
        source_extra.update({"paragraph": paragraph, "format": "docx"})
    else:
        if not lines:
            raise EvidenceError("文本或 Markdown 文档证据必须指定 --lines")
        text = data.decode("utf-8", errors="replace")
        excerpt, start, end = _select_lines(text, lines)
        locator = f"L{start}-L{end}"
        source_extra.update({"line_start": start, "line_end": end, "format": "text"})

    return add_evidence(
        case_dir,
        kind="document_span",
        title=title,
        source={
            "work_id": source_spec.work_id,
            "display_name": source_spec.display_name,
            "commit": locked_commit,
            "path": document_path,
            "locator": locator,
            "object_hash": blob_hash,
            **source_extra,
        },
        excerpt=excerpt,
        facts=[],
        raw_ref=raw_ref,
    )


def _numstat_rows(repo: Path, commit: str, paths: Iterable[str]) -> list[list[Any]]:
    args = ["show", "--format=", "--numstat", commit]
    path_list = list(paths)
    if path_list:
        args.extend(["--", *path_list])
    output = str(_run_git(repo, *args))
    rows: list[list[Any]] = []
    for line in output.splitlines():
        parts = line.split("\t", 2)
        if len(parts) == 3:
            rows.append([parts[2], parts[0], parts[1]])
    return rows


def capture_commit(
    case_dir: Path,
    *,
    title: str,
    commit: str,
    paths: Iterable[str] = (),
    work_id: str | None = None,
    repo: str | Path | None = None,
    display_name: str | None = None,
    external_id: str | None = None,
    works_path: Path = DEFAULT_WORKS,
) -> str:
    source_spec = _resolve_source(
        work_id=work_id,
        repo=repo,
        display_name=display_name,
        external_id=external_id,
        works_path=works_path,
    )
    locked_commit = _lock_commit(source_spec.repo, commit)
    fields = str(
        _run_git(
            source_spec.repo,
            "show",
            "-s",
            "--format=%H%x1f%P%x1f%an%x1f%aI%x1f%cn%x1f%cI%x1f%s%x1f%b",
            locked_commit,
        )
    ).rstrip("\n").split("\x1f", 7)
    if len(fields) != 8:
        raise EvidenceError("无法解析 commit 元数据")
    commit_hash, parents, author, author_time, committer, commit_time, subject, body = fields
    path_list = sorted(set(paths))
    rows = _numstat_rows(source_spec.repo, locked_commit, path_list)

    additions = sum(int(row[1]) for row in rows if str(row[1]).isdigit())
    deletions = sum(int(row[2]) for row in rows if str(row[2]).isdigit())
    facts = [
        {"label": "父提交", "value": parents.split() if parents else []},
        {"label": "作者", "value": author},
        {"label": "作者时间", "value": author_time},
        {"label": "提交者", "value": committer},
        {"label": "提交时间", "value": commit_time},
        {"label": "文件数", "value": len(rows)},
        {"label": "新增行", "value": additions},
        {"label": "删除行", "value": deletions},
    ]
    locator = "commit" if not path_list else "paths: " + ", ".join(path_list)
    excerpt = subject if not body.strip() else f"{subject}\n\n{body.strip()}"
    return add_evidence(
        case_dir,
        kind="git_commit",
        title=title,
        source={
            "work_id": source_spec.work_id,
            "display_name": source_spec.display_name,
            "commit": commit_hash,
            "path": ", ".join(path_list),
            "locator": locator,
            "object_hash": commit_hash,
        },
        excerpt=excerpt,
        facts=facts,
        table={"columns": ["文件", "新增", "删除"], "rows": rows},
    )


def _comparison_summary(data: dict[str, Any]) -> tuple[dict[str, Any], str, list[dict[str, Any]], dict[str, Any]]:
    if data.get("schema") != "review_case.commit_pair.v1":
        raise EvidenceError("comparison 目前只接受 compare-commits 生成的 commit_pair.v1")
    left = data.get("left", {})
    right = data.get("right", {})
    blob = data.get("blob_overlap", {})
    ast = data.get("ast_overlap") or {}
    source = {
        "work_id": f"{left.get('work_id', '')} ↔ {right.get('work_id', '')}",
        "display_name": f"{left.get('display_name', left.get('work_id', '左侧'))} ↔ "
        f"{right.get('display_name', right.get('work_id', '右侧'))}",
        "commit": f"{left.get('commit', '')} ↔ {right.get('commit', '')}",
        "path": "",
        "locator": "commit pair",
        "object_hash": "",
    }
    facts = [
        {"label": "相同内容实例", "value": blob.get("content_instances", {}).get("shared_instances", 0)},
        {"label": "同路径实例", "value": blob.get("path_aligned_instances", {}).get("shared_instances", 0)},
        {"label": "跨路径实例", "value": blob.get("relocated_instances", 0)},
        {"label": "共享唯一 Blob", "value": blob.get("shared_unique_blobs", 0)},
    ]
    if ast:
        facts.extend(
            [
                {"label": "AST 相同结构实例", "value": ast.get("shape_instances", {}).get("shared_instances", 0)},
                {"label": "AST 改名或移动实例", "value": ast.get("renamed_or_relocated_instances", 0)},
            ]
        )

    rows: list[list[Any]] = []
    for item in blob.get("relocated_examples", [])[:20]:
        rows.append(
            [
                "Blob 跨路径",
                ", ".join(item.get("left_paths", [])),
                ", ".join(item.get("right_paths", [])),
                item.get("blob", "")[:12],
            ]
        )
    for item in ast.get("examples", [])[:20]:
        left_item = item.get("left", {})
        right_item = item.get("right", {})
        rows.append(
            [
                "AST 结构",
                f"{left_item.get('path', '')}:{left_item.get('symbol', '')}",
                f"{right_item.get('path', '')}:{right_item.get('symbol', '')}",
                str(item.get("shape", ""))[:12],
            ]
        )
    excerpt = (
        f"比较 {left.get('commit', '')[:12]} 与 {right.get('commit', '')[:12]}："
        f"相同内容实例 {blob.get('content_instances', {}).get('shared_instances', 0)}，"
        f"同路径 {blob.get('path_aligned_instances', {}).get('shared_instances', 0)}，"
        f"跨路径 {blob.get('relocated_instances', 0)}。"
    )
    return source, excerpt, facts, {"columns": ["匹配类型", "左侧", "右侧", "指纹"], "rows": rows}


def capture_comparison(case_dir: Path, *, title: str, fact_file: Path) -> str:
    fact_file = fact_file.resolve()
    try:
        data = json.loads(fact_file.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise EvidenceError(f"无法读取比较结果 {fact_file}: {exc}") from exc
    source, excerpt, facts, table = _comparison_summary(data)
    sha256 = _sha256_file(fact_file)
    source["object_hash"] = sha256
    return add_evidence(
        case_dir,
        kind="fingerprint_comparison",
        title=title,
        source=source,
        excerpt=excerpt,
        facts=facts,
        table=table,
        raw_ref={"path": _display_path(fact_file), "sha256": sha256},
    )


def capture_search(
    case_dir: Path,
    *,
    title: str,
    commit: str,
    pattern: str,
    paths: Iterable[str] = (),
    work_id: str | None = None,
    repo: str | Path | None = None,
    display_name: str | None = None,
    external_id: str | None = None,
    works_path: Path = DEFAULT_WORKS,
) -> str:
    source_spec = _resolve_source(
        work_id=work_id,
        repo=repo,
        display_name=display_name,
        external_id=external_id,
        works_path=works_path,
    )
    locked_commit = _lock_commit(source_spec.repo, commit)
    path_list = sorted(set(paths))
    args = ["git", "-C", str(source_spec.repo), "grep", "-n", "-I", "-E", "-e", pattern, locked_commit]
    if path_list:
        args.extend(["--", *path_list])
    result = subprocess.run(args, check=False, capture_output=True, text=True)
    if result.returncode not in (0, 1):
        raise EvidenceError(result.stderr.strip() or "git grep 执行失败")
    matches = result.stdout.splitlines()
    rows: list[list[Any]] = []
    for line in matches[:100]:
        prefix = f"{locked_commit}:"
        cleaned = line[len(prefix) :] if line.startswith(prefix) else line
        parts = cleaned.split(":", 2)
        if len(parts) == 3:
            rows.append([f"{parts[0]}:{parts[1]}", parts[2]])
        else:
            rows.append([cleaned, ""])
    scope = ", ".join(path_list) if path_list else "整个提交"
    excerpt = f"在 {locked_commit[:12]} 的 {scope} 中检索 {pattern!r}，命中 {len(matches)} 处。"
    return add_evidence(
        case_dir,
        kind="search_result",
        title=title,
        source={
            "work_id": source_spec.work_id,
            "display_name": source_spec.display_name,
            "commit": locked_commit,
            "path": scope,
            "locator": f"git grep {pattern!r}",
            "object_hash": locked_commit,
        },
        excerpt=excerpt,
        facts=[
            {"label": "检索式", "value": pattern},
            {"label": "范围", "value": scope},
            {"label": "命中数", "value": len(matches)},
        ],
        table={"columns": ["位置", "命中内容"], "rows": rows},
    )


def capture_search_result(case_dir: Path, *, title: str, fact_file: Path) -> str:
    fact_file = fact_file.resolve()
    try:
        data = json.loads(fact_file.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise EvidenceError(f"无法读取检索结果 {fact_file}: {exc}") from exc
    schema = data.get("schema")
    target = data.get("target", {})
    rows: list[list[Any]] = []
    if schema == "review_case.search_candidates.v3":
        for item in data.get("candidates", [])[:50]:
            rows.append(
                [
                    item.get("display_name", item.get("work_id", "")),
                    str(item.get("commit", ""))[:12],
                    item.get("blob_overlap", {}).get("content_instances", {}).get("balanced", 0),
                    item.get("ast_overlap", {}).get("shape_instances", {}).get("balanced", 0),
                ]
            )
        columns = ["候选", "Commit", "Blob", "AST"]
        query = "HEAD Blob/AST 候选检索"
    elif schema == "review_case.history_blob_candidates.v1":
        for item in data.get("candidates", [])[:50]:
            rows.append(
                [
                    item.get("display_name", item.get("work_id", "")),
                    str(item.get("history_head", ""))[:12],
                    item.get("target_history_containment", 0),
                    item.get("shared_unique_blobs", 0),
                ]
            )
        columns = ["候选", "历史 Commit", "目标覆盖率", "共享 Blob"]
        query = "历史 Commit Blob 检索"
    else:
        raise EvidenceError(f"search --file 不支持 schema: {schema}")

    sha256 = _sha256_file(fact_file)
    if schema == "review_case.search_candidates.v3":
        target = {
            "work_id": data.get("target_work_id", ""),
            "display_name": data.get("target_display_name", data.get("target_work_id", "")),
            "commit": data.get("target_commit", ""),
        }
    source = {
        "work_id": str(target.get("work_id", "search")),
        "display_name": str(target.get("display_name", target.get("work_id", "检索结果"))),
        "commit": str(target.get("commit", "")),
        "path": "",
        "locator": query,
        "object_hash": sha256,
    }
    excerpt = f"{query}返回 {len(data.get('candidates', []))} 个候选。"
    return add_evidence(
        case_dir,
        kind="search_result",
        title=title,
        source=source,
        excerpt=excerpt,
        facts=[
            {"label": "检索类型", "value": query},
            {"label": "候选数", "value": len(data.get("candidates", []))},
        ],
        table={"columns": columns, "rows": rows},
        raw_ref={"path": _display_path(fact_file), "sha256": sha256},
    )
